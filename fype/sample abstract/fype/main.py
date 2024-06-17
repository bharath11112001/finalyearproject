from flask import Flask, render_template, redirect, request, session, url_for
from datetime import datetime
import datetime
import os
from werkzeug.utils import secure_filename
from flask import send_from_directory, abort
from docx import Document
import mysql.connector
from docx.shared import Inches
from docx2pdf import convert
import pythoncom
import time
import base64
from io import BytesIO
import io
import fitz  # PyMuPDF
from fpdf import FPDF
from flask import send_file
from flask import Response
from flask_socketio import SocketIO, emit, join_room, leave_room
from engineio.payload import Payload
Payload.max_decode_packets = 200
from werkzeug.utils import secure_filename


app = Flask(__name__, static_url_path='/static')
app.secret_key = 'abcdef'


socketio = SocketIO(app)


_users_in_room = {} # stores room wise user list
_room_of_sid = {} # stores room joined by an used
_name_of_sid = {} # stores display name of users


mydb = mysql.connector.connect(
    host="localhost",
    user="root",
    password="",
    charset="utf8",
    use_pure=True,
    database="fyp_db"
)

@app.route('/',methods=['POST','GET'])
def index():

    return render_template('index.html')

@app.route('/login',methods=['POST','GET'])
def login():
    
    
    msg=""
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        cursor = mydb.cursor()
        cursor.execute('SELECT * FROM fyp_admin WHERE username = %s AND password = %s', (username, password))
        account = cursor.fetchone()
        
        if account:
            session['username'] = username
            session['user_type'] = 'admin'
            msg="success"
        else:
            msg="fail"
        

    return render_template('login.html',msg=msg)



@app.route('/hod_forget',methods=['POST','GET'])
def hod_forget():

    msg=""
    mess=""
    st=""
    email=""
    password=""
    if request.method=='POST':
        username=request.form['username']
        password=request.form['password']

        cursor=mydb.cursor()
        cursor.execute(" UPDATE fyp_staff SET password=%s WHERE staff_id=%s ", (password, username))
        mydb.commit()
        cursor.close()
        msg="success"
        cursor = mydb.cursor()
        cursor.execute("SELECT * FROM fyp_staff where staff_id=%s", (username, ))
        data2 = cursor.fetchone()
        cursor.close()
        email=data2[3]
        password=data2[6]
        st="1"
        mess=f"Reminder: Hi , your password changed successfully and the new password is {password}"
    else:
        msg="fail"
        
    return render_template('hod_forget.html', msg=msg, email=email, password=password, st=st, mess=mess)


@app.route('/staff_forget',methods=['POST','GET'])
def staff_forget():

    msg=""
    mess=""
    st=""
    email=""
    password=""
    if request.method=='POST':
        username=request.form['username']
        password=request.form['password']

        cursor=mydb.cursor()
        cursor.execute(" UPDATE fyp_staff SET password=%s WHERE staff_id=%s ", (password, username))
        mydb.commit()
        cursor.close()
        msg="success"
        cursor = mydb.cursor()
        cursor.execute("SELECT * FROM fyp_staff where staff_id=%s", (username, ))
        data2 = cursor.fetchone()
        cursor.close()
        email=data2[3]
        password=data2[6]
        st="1"
        mess=f"Reminder: Hi , your password changed successfully and the new password is {password}"
    else:
        msg="fail"
        
    return render_template('staff_forget.html', msg=msg, email=email, password=password, st=st, mess=mess)

@app.route('/stu_forget',methods=['POST','GET'])
def stu_forget():

    msg=""
    mess=""
    st=""
    email=""
    password=""
    if request.method=='POST':
        username=request.form['username']
        password=request.form['password']

        cursor=mydb.cursor()
        cursor.execute(" UPDATE fyp_student SET password=%s WHERE reg_no=%s ", (password, username))
        mydb.commit()
        cursor.close()
        msg="success"
        cursor = mydb.cursor()
        cursor.execute("SELECT * FROM fyp_student where reg_no=%s", (username, ))
        data2 = cursor.fetchone()
        cursor.close()
        email=data2[6]
        password=data2[11]
        st="1"
        mess=f"Reminder: Hi , your password changed successfully and the new password is {password}"
    else:
        msg="fail"
        
    return render_template('stu_forget.html', msg=msg, email=email, password=password, st=st, mess=mess)


@app.route('/addstu',methods=['POST','GET'])
def addstu():
    if 'username' not in session or session.get('user_type') != 'admin':
        print("Please log in as a hod to access the page.", 'danger')
        return redirect(url_for('login'))
    msg=""
    st=""
    name=""
    email=""
    mess=""
    reg_no=""
    password=""
    if request.method=='POST':
        reg_no=request.form['reg_no']
        name=request.form['name']
        gender=request.form['gender']
        dob=request.form['dob']
        mobile=request.form['mobile']
        captcha=request.form['captcha']
        email=request.form['email']
        address=request.form['address']
        dept=request.form['dept']
        batch_year=request.form['batch_year']
        password=request.form['password']

        
        now = datetime.datetime.now()
        date_join=now.strftime("%d-%m-%Y")
        mycursor = mydb.cursor()

        mycursor.execute("SELECT count(*) FROM fyp_student where reg_no=%s",(reg_no, ))
        cnt = mycursor.fetchone()[0]
        if cnt==0:
            mycursor.execute("SELECT max(id)+1 FROM fyp_student")
            maxid = mycursor.fetchone()[0]
            if maxid is None:
                maxid=1
            sql = "INSERT INTO fyp_student(id, name, reg_no, gender, dob, mobile, email, address,  dept, batch_year, password,date_join, captcha) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,%s)"
            val = (maxid, name, reg_no, gender, dob, mobile, email, address,  dept, batch_year, password,date_join, captcha)
            mycursor.execute(sql, val)
            mydb.commit()

            msg="success"
            st="1"
            mess = f"Reminder: Hi {name}, your username is {reg_no} and password is {password}!"
            mycursor.close()
        else:
            msg="fail"
            
    return render_template('addstu.html',msg=msg, email=email, mess=mess, st=st, reg_no=reg_no, password=password, name=name)


@app.route('/addstaff',methods=['POST','GET'])
def addstaff():
    if 'username' not in session or session.get('user_type') != 'admin':
        print("Please log in as a hod to access the page.", 'danger')
        return redirect(url_for('login'))
    msg=""
    st=""
    name=""
    email=""
    mess=""
    staff_id=""
    password=""
    if request.method=='POST':
        
        staff_id=request.form['staff_id']
        name=request.form['name']       
        mobile=request.form['mobile']
        email=request.form['email']
        captcha=request.form['captcha']
        dept=request.form['dept']
        location=request.form['location']        
        password=request.form['password']
        staff_type=request.form['staff_type']
        
       

        
        now = datetime.datetime.now()
        date_join=now.strftime("%Y-%m-%d")
        mycursor = mydb.cursor()

        mycursor.execute("SELECT count(*) FROM fyp_staff where staff_id=%s",(staff_id, ))
        cnt = mycursor.fetchone()[0]
        if cnt==0:
            mycursor.execute("SELECT max(id)+1 FROM fyp_staff")
            maxid = mycursor.fetchone()[0]
            if maxid is None:
                maxid=1
            sql = "INSERT INTO fyp_staff(id, name, mobile, email, location,  staff_id, password, staff_type,date_join, captcha, dept) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s,%s,%s)"
            val = (maxid, name, mobile, email, location,  staff_id, password, staff_type,date_join, captcha, dept)
            
            
            mycursor.execute(sql, val)
            mydb.commit()

            msg="success"
            st="1"
            mess = f"Reminder: Hi {name}, your username is {staff_id} and password is {password}!"
            mycursor.close()
        else:
            msg="fail"
            
    return render_template('addstaff.html',msg=msg, email=email, mess=mess, st=st, staff_id=staff_id, password=password, name=name)


@app.route('/profile',methods=['POST','GET'])
def profile():
    if 'username' not in session or session.get('user_type') != 'admin':
        print("Please log in as a hod to access the page.", 'danger')
        return redirect(url_for('login'))
    
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM fyp_student")
    value = mycursor.fetchall()
    mycursor.close()

    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM fyp_staff order by id")
    value2 = mycursor.fetchall()
    mycursor.close()
    if request.method == 'POST':
        search = request.form.get('search')
        cursor = mydb.cursor()
        cursor.execute("SELECT * FROM fyp_student WHERE dept LIKE %s", ('%' + search + '%',))
        value = cursor.fetchall()
        cursor.execute("SELECT * FROM fyp_staff WHERE dept LIKE %s", ('%' + search + '%',))
        value2 = cursor.fetchall()
        
        cursor.close()

    return render_template('profile.html', data=value, dataa=value2)


@app.route('/deletee/<string:id_data>', methods=['GET'])
def deletee(id_data):
    cursor=mydb.cursor()
    cursor.execute("DELETE FROM fyp_student WHERE id=%s", (id_data,))
    mydb.commit()
    print("Record Has Been Deleted Successfully")
    return redirect(url_for('profile'))


@app.route('/delete/<string:id_data>', methods=['GET'])
def delete(id_data):
    cursor=mydb.cursor()
    cursor.execute("DELETE FROM fyp_staff WHERE id=%s", (id_data,))
    mydb.commit()
    print("Record Has Been Deleted Successfully")
    return redirect(url_for('profile'))



@app.route('/hodlog',methods=['POST','GET'])
def hodlog():   

    msg=""
    if request.method == 'POST':
        staff_id = request.form['staff_id']
        password = request.form['password']
       
        
        cursor = mydb.cursor()
        cursor.execute('SELECT * FROM fyp_staff WHERE staff_type="hod" AND staff_id = %s AND password = %s', (staff_id, password))
        account = cursor.fetchone()
        
        if account:
            session['staff_id'] = staff_id
            session['user_type'] = 'hod'
            msg="success"
            
        else:
            msg="fail"
            
    return render_template('hodlog.html', msg=msg)


@app.route('/viewhod',methods=['POST','GET'])
def viewhod():
    
    if 'staff_id' not in session or session.get('user_type') != 'hod':
        print("Please log in as a hod to access the page.", 'danger')
        return redirect(url_for('hodlog'))
    data1=""
    data3=""
    msg=""

    staff_id = session.get('staff_id')
    cursor = mydb.cursor()
    cursor.execute("SELECT staff_type, dept FROM fyp_staff WHERE staff_id = %s", (staff_id,))
    user_data = cursor.fetchone()
    user_type, user_dept = user_data[0], user_data[1]

    

    if user_type == 'hod':
        
        
        cursor.execute("SELECT * FROM fyp_staff WHERE staff_type='staff' AND dept = %s", (user_dept,))
        data1 = cursor.fetchall()
        cursor.execute("SELECT * FROM fyp_student WHERE dept = %s", (user_dept,))
        data3 = cursor.fetchall()
        

    cursor.execute("SELECT DISTINCT reg_no FROM fyp_staff_allocation")
    allocated_reg_nos = [row[0] for row in cursor.fetchall()]
    cursor.close()

    staff_id = request.args.get('staff_id')
    cursor = mydb.cursor()
    cursor.execute("SELECT * FROM fyp_staff_allocation WHERE staff_id = %s", (staff_id,))
    data2 = cursor.fetchall()
    cursor.close()
        
    

    if request.method == 'POST':
        selected_staff_id = request.form.get('staff')
        selected_students = request.form.getlist('selected_staff[]')
        
        

        
        cursor = mydb.cursor(buffered=True)
        cursor.execute("SELECT staff_id, dept FROM fyp_staff WHERE id = %s", (selected_staff_id,))
        staff_details = cursor.fetchone()
        staff_id = staff_details[0]
        department = staff_details[1]

        
        student_details = {}
        for student_id in selected_students:
            cursor.execute("SELECT reg_no, dept, name FROM fyp_student WHERE id = %s", (student_id,))
            student_info = cursor.fetchone()
            if student_info:
                reg_no = student_info[0]
                dept = student_info[1]
                name = student_info[2]
                student_details[student_id] = {'reg_no': reg_no, 'dept': dept, 'name': name}

        
        for student_id, details in student_details.items():
            reg_no = details['reg_no']
            dept = details['dept']
            name = details['name']
            cursor = mydb.cursor()

            
            cursor.execute("SELECT max(id)+1 FROM fyp_staff_allocation")
            maxid = cursor.fetchone()[0]
            if maxid is None:
                maxid=1
            cursor.execute("INSERT INTO fyp_staff_allocation (id, staff_id, dept, reg_no, name) VALUES (%s, %s, %s, %s, %s)", (maxid, staff_id, dept, reg_no, name))

        mydb.commit()
        msg="success"

        cursor.close()
        

    return render_template('viewhod.html', fyp_staff=data1, fyp_student=data3, allocated_reg_nos=allocated_reg_nos, fyp_staff_allocation=data2, staff_id=staff_id, msg=msg)


@app.route('/allocatedet',methods=['POST','GET'])
def allocatedet():

    if 'staff_id' not in session or session.get('user_type') != 'hod':
        print("Please log in as a hod to access the page.", 'danger')
        return redirect(url_for('hodlog'))
    
    data1=""
    data3=""

    staff_id = session.get('staff_id')
    cursor = mydb.cursor()
    cursor.execute("SELECT staff_type, dept FROM fyp_staff WHERE staff_id = %s", (staff_id,))
    user_data = cursor.fetchone()
    user_type, user_dept = user_data[0], user_data[1]

    

    if user_type == 'hod':
        
        
        cursor.execute("SELECT * FROM fyp_staff WHERE staff_type='staff' AND dept = %s", (user_dept,))
        data1 = cursor.fetchall()
        cursor.execute("SELECT * FROM fyp_student WHERE dept = %s", (user_dept,))
        data3 = cursor.fetchall()
        

    cursor.execute("SELECT DISTINCT reg_no FROM fyp_staff_allocation")
    allocated_reg_nos = [row[0] for row in cursor.fetchall()]
    cursor.close()

    staff_id = request.args.get('staff_id')
    cursor = mydb.cursor()
    cursor.execute("SELECT * FROM fyp_staff_allocation WHERE staff_id = %s", (staff_id,))
    data2 = cursor.fetchall()
    cursor.close()
        
    

    if request.method == 'POST':
        selected_staff_id = request.form.get('staff')
        selected_students = request.form.getlist('selected_staff[]')
        name = request.form.get('name')
        sname = request.form.get('sname')

        
        cursor = mydb.cursor()
        cursor.execute("SELECT staff_id, dept FROM fyp_staff WHERE id = %s", (selected_staff_id,))
        staff_details = cursor.fetchone()
        staff_id = staff_details[0]
        department = staff_details[1]

        
        student_details = {}
        for student_id in selected_students:
            cursor.execute("SELECT reg_no, dept FROM fyp_student WHERE id = %s", (student_id,))
            student_info = cursor.fetchone()
            if student_info:
                reg_no = student_info[0]
                dept = student_info[1]
                student_details[student_id] = {'reg_no': reg_no, 'dept': dept}

        
        for student_id, details in student_details.items():
            reg_no = details['reg_no']
            dept = details['dept']
            cursor = mydb.cursor()

            
            cursor.execute("SELECT max(id)+1 FROM fyp_staff_allocation")
            maxid = cursor.fetchone()[0]
            if maxid is None:
                maxid=1
            cursor.execute("INSERT INTO fyp_staff_allocation (id, staff_id, dept, reg_no, name, sname) VALUES (%s, %s, %s, %s, %s, %s)", (maxid, staff_id, dept, reg_no, name, sname))

        mydb.commit()

        cursor.close()
        return redirect(url_for('hodlog'))

    return render_template('allocatedet.html', fyp_staff=data1, fyp_student=data3, allocated_reg_nos=allocated_reg_nos, fyp_staff_allocation=data2, staff_id=staff_id)

    

@app.route('/stafflog',methods=['POST','GET'])
def stafflog():    

    msg=""
    if request.method == 'POST':
        staff_id = request.form['staff_id']
        password = request.form['password']
        
        
        cursor = mydb.cursor()
        cursor.execute('SELECT * FROM fyp_staff WHERE staff_id = %s AND password = %s', (staff_id, password))
        account = cursor.fetchone()
        
        if account:
            session['staff_id'] = staff_id
            session['user_type'] = 'staff'
            msg="success"
            
        else:
            msg="fail"
            
    return render_template('stafflog.html', msg=msg)




@app.route('/staffview', methods=['POST', 'GET'])
def staffview():
    if 'staff_id' not in session or session.get('user_type') != 'staff':
        print("Please log in as a admin to access the page.", 'danger')
        return redirect(url_for('stafflog'))
    
    cursor = mydb.cursor()
    staff_id = session.get('staff_id')
    cursor.execute("SELECT * FROM fyp_staff_allocation WHERE staff_id = %s", (staff_id,))
    data = cursor.fetchall()
    cursor.close()

    return render_template('staffview.html', fyp_staff_allocation=data)

def allowed_file(filename):
    ALLOWED_EXTENSIONS = {'jpg', 'jpeg', 'png'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route('/viewabs', methods=['POST', 'GET'])
def viewabs():
    if 'staff_id' not in session or session.get('user_type') != 'staff':
        print("Please log in as a hod to access the page.", 'danger')
        return redirect(url_for('stafflog'))

    cursor = mydb.cursor()
    staff_id = session.get('staff_id')
    cursor.execute("SELECT * FROM fyp_staff_allocation WHERE staff_id = %s", (staff_id,))
    data1 = cursor.fetchall()
    cursor.close()

    if request.method == 'POST':
        dept = request.form.get('dept')
        staff_id = request.form.get('staff_id')

        if 'image' in request.files:
            image = request.files['image']

            if image and allowed_file(image.filename):
                filename = secure_filename(image.filename)
                image_path = 'E:/fyp/static/sign/' + filename
                image.save(image_path)

                mycursor = mydb.cursor()
                mycursor.execute("SELECT max(id)+1 FROM fyp_digital_signature")
                maxid = mycursor.fetchone()[0]
                if maxid is None:
                    maxid = 1

                    # Store each file path separately in a new row

                sql = "INSERT INTO fyp_digital_signature (id, dept, staff_id, signature_link) VALUES (%s, %s, %s, %s)"
                val = (maxid, dept, staff_id, filename)
                mycursor.execute(sql, val)
                mydb.commit()
                mycursor.close()
                print("successful")
            else:
                print("unsuccessful")
        else:
            print("Invalid file format. Allowed formats: jpg, jpeg, png")
    else:
        print("No file uploaded")
    
        
    staff_id = session.get('staff_id')
    cursor = mydb.cursor(buffered=True)
    
    try:
        reg_no = request.args.get('reg_no')
        cursor.execute("SELECT * FROM fyp_title WHERE reg_no = %s", (reg_no,))
        data = cursor.fetchall()

        cursor.execute("SELECT * FROM fyp_title WHERE staff_id = %s", (staff_id,))
        user_dataa = cursor.fetchone()

        cursor.execute("SELECT reg_no FROM fyp_title WHERE staff_id = %s", (staff_id,))
        user_data = cursor.fetchone()
        user_name = user_data[0]
        print(user_name)

        cursor.execute("SELECT * FROM fyp_student WHERE reg_no = %s", (user_name,))
        student_data = cursor.fetchone()

    finally:
        cursor.close()

    return render_template('viewabs.html', fyp_title=data, user_name=user_name, student_data=student_data, user_dataa=user_dataa, fyp_staff_allocation=data1)

@app.route('/download_document/<path:doc_path>', methods=['GET'])
def download_document(doc_path):
    
    uploads_folder = 'E:/fyp/static/uploads/'  # Adjust the folder structure as needed
    return send_from_directory(uploads_folder, doc_path, as_attachment=True)


@app.route('/downloadd/<path:doc_path>', methods=['GET'])
def downloadd(doc_path):
    
    uploads_folder = 'E:/fyp/static/report/'  # Adjust the folder structure as needed
    return send_from_directory(uploads_folder, doc_path, as_attachment=True)


@app.route('/delete_digital_signature/<reg_no>', methods=['GET'])
def delete_digital_signature(reg_no):
    if 'staff_id' not in session or session.get('user_type') != 'staff':
        print("Please log in as a hod to access the page.", 'danger')
        return redirect(url_for('stafflog'))

    cursor = mydb.cursor()

    try:
        # Update the digital_signature_link to NULL for the specified reg_no
        update_query = "UPDATE fyp_title SET digital_signature_link = NULL WHERE reg_no = %s"
        cursor.execute(update_query, (reg_no,))
        mydb.commit()
        print(f"Digital signature link for reg_no {reg_no} deleted successfully.")
    except Exception as e:
        print(f"Error deleting digital signature link for reg_no {reg_no}: {e}")
        mydb.rollback()
    finally:
        cursor.close()

    return redirect(url_for('viewabs'))

@app.route('/reviewupdate',methods=['POST','GET'])
def reviewupdate():
    
    if 'staff_id' not in session or session.get('user_type') != 'staff':
        print("Please log in as a hod to access the page.", 'danger')
        return redirect(url_for('stafflog'))

    staff_id = session.get('staff_id')
    cursor = mydb.cursor()
    cursor.execute("SELECT * FROM fyp_staff WHERE staff_id = %s", (staff_id,))
    data = cursor.fetchone()
    cursor.execute("SELECT * FROM fyp_review WHERE staff_id = %s", (staff_id,))
    data1 = cursor.fetchall()
    cursor.close()
    
    msg=""
    if request.method=='POST':
        staff_id=request.form.get('staff_id')
        dept=request.form.get('dept')
        name=request.form.get('name')
        review_date=request.form['review_date']
        review_time=request.form['review_time']
        requirement=request.form['requirement']        

        mycursor = mydb.cursor()
        mycursor.execute("SELECT max(id)+1 FROM fyp_review")
        maxid = mycursor.fetchone()[0]
        if maxid is None:
            maxid=1
        sql = "INSERT INTO fyp_review(id, staff_id, dept, review_date, review_time, requirement, name) VALUES (%s, %s, %s, %s, %s, %s, %s)"
        val = (maxid, staff_id, dept, review_date, review_time, requirement, name)
        mycursor.execute(sql, val)
        mydb.commit()
        msg="success"
            
    

    return render_template('reviewupdate.html', data=data, msg=msg, fyp_review=data1)

@app.route('/edit', methods=['GET', 'POST'])
def edit():
    if request.method=='POST':
        id_data=request.form['id']
        date=request.form['date']
        time=request.form['time']
        remark=request.form['remark']

        cursor=mydb.cursor()
        cursor.execute(" UPDATE fyp_review SET review_date=%s, review_time=%s, requirement=%s WHERE id=%s ", (date, time, remark, id_data))
        mydb.commit()
        cursor.close()
        print("Data Updated Successfully")
    
    return redirect(url_for('reviewupdate'))


@app.route('/viewreport', methods=['POST', 'GET'])
def viewreport():
    
    if 'staff_id' not in session or session.get('user_type') != 'staff':
        print("Please log in as a admin to access the page.", 'danger')
        return redirect(url_for('stafflog'))
    
    cursor = mydb.cursor()
    staff_id = session.get('staff_id')
    cursor.execute("SELECT * FROM fyp_staff_allocation WHERE staff_id = %s", (staff_id,))
    data = cursor.fetchall()
    cursor.close()

    
    reg_no = request.args.get('reg_no')
    cursor = mydb.cursor()
    cursor.execute("SELECT * FROM fyp_review_report WHERE reg_no = %s", (reg_no,))
    reports_data = cursor.fetchall()
    cursor.close()

    return render_template('viewreport.html', fyp_staff_allocation=data, reg_no=reg_no, fyp_review_report=reports_data)

@app.route('/download_report/<int:review_id>', methods=['GET'])
def download_report(review_id):
    try:
        uploads_folder = 'E:/fyp/static/report/'

        cursor = mydb.cursor()
        cursor.execute("SELECT * FROM fyp_review_report WHERE review_id = %s", (review_id,))
        review_data = cursor.fetchone()
        cursor.close()

        if review_data:
           
            file_names = review_data[5].split(', ')  
            file_index = request.args.get('file_index', default=0, type=int)
            if 0 <= file_index < len(file_names):
                selected_file = file_names[file_index]
                file_path = os.path.join(uploads_folder, selected_file)
                return send_from_directory(uploads_folder, selected_file, as_attachment=True)
            else:
                abort(404)  
        else:
            abort(404) 
    except FileNotFoundError:
        abort(404) 
    except Exception as e:
        print("Error:", str(e))
        abort(500)
        
@app.route('/result', methods=['POST', 'GET'])
def result():
    
    if 'staff_id' not in session or session.get('user_type') != 'staff':
        print("Please log in as a admin to access the page.", 'danger')
        return redirect(url_for('stafflog'))

    if request.method == 'POST':
        titlee = request.form.get('titlee')
        reg_no = request.form.get('reg_no')
        dept = request.form.get('dept')
        staff_id = request.form.get('staff_id')
        marks = request.form.get('marks')
        remark = request.form.get('remark')

        mycursor = mydb.cursor()
        mycursor.execute("SELECT max(id)+1 FROM fyp_review_result")
        maxid = mycursor.fetchone()[0]
        if maxid is None:
            maxid=1
        sql = "INSERT INTO fyp_review_result(id, reg_no, titlee, dept, marks, remark, staff_id) VALUES (%s, %s, %s, %s, %s, %s, %s)"
        val = (maxid, reg_no, titlee, dept, marks, remark, staff_id)
        mycursor.execute(sql, val)
        mydb.commit()
        mycursor.close()
        msg="success"
    
    cursor = mydb.cursor()
    staff_id = session.get('staff_id')
    cursor.execute("SELECT * FROM fyp_staff_allocation WHERE staff_id = %s", (staff_id,))
    data = cursor.fetchall()
    cursor.close()

    
    reg_no = request.args.get('reg_no')
    cursor = mydb.cursor()
    cursor.execute("SELECT * FROM fyp_title WHERE reg_no = %s", (reg_no,))
    reports_data = cursor.fetchall()
    cursor.close()

    return render_template('result.html', fyp_staff_allocation=data, reg_no=reg_no, fyp_title=reports_data)
        
@app.route('/viewreview', methods=['POST', 'GET'])
def viewreview():
    msg = ""
    mobile=""
    st=""
    review_id=""
    reg_no=""
    mess=""
    staff_id=""
    if 'reg_no' not in session or session.get('user_type') != 'student':
        print("Please log in as a student to access the page.", 'danger')
        return redirect(url_for('stulog'))

    reg_no = session.get('reg_no')

    if request.method == 'POST':
        review_id = request.form.get('review_id')
        date = request.form.get('date')
        staff_id = request.form.get('Staff_id')
        dept = request.form.get('dept')
        pdf_files = request.files.getlist('pdf[]')

        file_paths = []

        if pdf_files:
            for pdf in pdf_files:  # Iterate through the list of files
                doc_filename = secure_filename(pdf.filename)
                pdf.save('E:/fyp/static/report/' + doc_filename)
                doc_path = doc_filename  # Adjust the folder structure as needed
                file_paths.append(doc_path)

            mycursor = mydb.cursor()
            
            mycursor.execute("SELECT max(id)+1 FROM fyp_review_report")
            maxid = mycursor.fetchone()[0]
            if maxid is None:
                maxid = 1

            # Convert the file_paths list to a string or tuple, depending on your database structure
            file_paths_str = ', '.join(file_paths)
            # or file_paths_tuple = tuple(file_paths)

            sql = "INSERT INTO fyp_review_report(id, review_id, date, staff_id, dept, upload_file, reg_no) VALUES (%s, %s, %s, %s, %s, %s, %s)"
            val = (maxid, review_id, date, staff_id, dept, file_paths_str, reg_no)
            mycursor.execute(sql, val)
            mydb.commit()

            msg = "success"
            st="1"
            mess = f"Reminder: Hi, register no {reg_no} has upload a file for {review_id} review!"
            mycursor.execute('SELECT mobile FROM fyp_staff WHERE staff_id = %s', (staff_id,))
            result = mycursor.fetchone()
            if result:
                mobile = result[0]
                
        else:
            msg = "fail"
        

    cursor = mydb.cursor()
    cursor.execute("SELECT staff_id FROM fyp_staff_allocation WHERE reg_no = %s", (reg_no,))
    user_data = cursor.fetchone()
    user_id = user_data[0]
    cursor.execute("select * from fyp_review WHERE staff_id = %s", (user_id,))
    data = cursor.fetchall()
    cursor.execute("select * from fyp_review_result WHERE reg_no = %s", (reg_no,))
    data1 = cursor.fetchall()
    cursor.close()
    

    return render_template('viewreview.html', fyp_review=data, user_id=user_id, msg=msg, fyp_review_result=data1, staff_id=staff_id, mobile=mobile, st=st, mess=mess, review_id=review_id, reg_no=reg_no)


@app.route('/stulog',methods=['POST','GET'])
def stulog():    

    msg=""
    if request.method == 'POST':
        reg_no = request.form['reg_no']
        password = request.form['password']
    
        
        cursor = mydb.cursor()
        cursor.execute('SELECT * FROM fyp_student WHERE reg_no = %s AND password = %s', (reg_no, password))
        account = cursor.fetchone()
        
        if account:
            session['reg_no'] = reg_no
            session['user_type'] = 'student'
            msg="success"
            
        else:
            msg="fail"
            
    return render_template('stulog.html', msg=msg)



@app.route('/abstract', methods=['POST', 'GET'])
def abstract():
    msg = ""
    msg_type = ""
    if 'reg_no' not in session or session.get('user_type') != 'student':
        print("Please log in as a student to access the page.", 'danger')
        return redirect(url_for('stulog'))

    if request.method == 'POST':
        dept = request.form.get('dept')
        reg_no = request.form.get('reg_no')
        staff_id = request.form.get('staff_id')
        name = request.form.get('name')
        doc_file = request.files.get('doc')
        if doc_file:
            doc_filename = secure_filename(doc_file.filename)
            doc_file.save('E:/fyp/static/uploads/' + doc_filename)
            doc_path = (doc_filename)  # Adjust the folder structure as needed

            # Check if a document with the same title already exists
            title = extract_title_from_doc(doc_file)
            if not is_document_title_duplicate(title):
                mycursor = mydb.cursor()
                mycursor.execute("SELECT MAX(id)+1 FROM fyp_title")
                maxid = mycursor.fetchone()[0]
                if maxid is None:
                    maxid = 1

                sql = "INSERT INTO fyp_title(id, dept, reg_no, staff_id, name, doc_path, titlee) VALUES (%s, %s, %s, %s, %s, %s, %s)"
                val = (maxid, dept, reg_no, staff_id, name, doc_path, title)
                mycursor.execute(sql, val)
                mydb.commit()
                msg = "Document uploaded successfully"
                msg_type = "success"
            else: 
                msg = "Title already exists!"
                msg_type = "danger"

    reg_no = session.get('reg_no')
    cursor = mydb.cursor()
    cursor.execute("SELECT * FROM fyp_staff_allocation WHERE reg_no = %s", (reg_no,))
    dat = cursor.fetchone()
    cursor.close()

    sta_id=dat[2]

    cursor = mydb.cursor()
    cursor.execute("SELECT * FROM fyp_staff WHERE staff_id = %s", (sta_id,))
    datt = cursor.fetchone()
    cursor.close()

    sname=datt[1]
                

    reg_no = session.get('reg_no')
    cursor = mydb.cursor()
    cursor.execute("SELECT * FROM fyp_staff_allocation WHERE reg_no = %s", (reg_no,))
    data1 = cursor.fetchall()
    cursor.close()

    # Retrieve the digital signature link based on reg_no
    
    cursor = mydb.cursor()
    cursor.execute("SELECT * FROM fyp_title WHERE reg_no = %s", (reg_no,))
    data = cursor.fetchall()
    cursor.close()


    return render_template('abstract.html', msg=msg, msg_type=msg_type, data1=data1, fyp_title=data, sname=sname)


def extract_title_from_doc(doc_file):
    doc = Document(doc_file)
    title = None
    for paragraph in doc.paragraphs:
        if paragraph.style.name == 'Title':
            title = paragraph.text
            break
    return title


def is_document_title_duplicate(title):
    cursor = mydb.cursor()
    cursor.execute("SELECT COUNT(*) FROM fyp_title WHERE titlee = %s", (title,))
    count = cursor.fetchone()[0]
    cursor.close()
    return count > 0


@app.route('/download_pdf/<path:digital_signature_link>', methods=['GET'])
def download_pdf(digital_signature_link):
    
    uploads_folder = 'E:/fyp/static/modified/'  # Adjust the folder structure as needed
    return send_from_directory(uploads_folder, digital_signature_link, as_attachment=True)


@app.route('/download_img', methods=['POST', 'GET'])
def download_img():
    if 'staff_id' not in session or session.get('user_type') != 'staff':
        print("Please log in as a hod to access the page.", 'danger')
        return redirect(url_for('stafflog'))

    msg=""

    staff_id = request.args.get('staff_id')

    cursor = mydb.cursor(buffered=True)
    cursor.execute("SELECT * FROM fyp_title WHERE id = %s", (staff_id,))
    data1 = cursor.fetchall()

    staffid = session.get('staff_id')
    cursor.execute("SELECT * FROM fyp_digital_signature WHERE staff_id = %s", (staffid,))
    data2 = cursor.fetchall()

    if data1 and data2:
        document_path = data1[0][5]
        signature_path = data2[0][3]

        # Initialize COM
        pythoncom.CoInitialize()

        try:
            # Open the Word document
            doc = Document(os.path.join('E:/fyp/static/uploads/', document_path))

            # Add the signature image to the Word document
            doc.add_picture(os.path.join('E:/fyp/static/sign/', signature_path), width=Inches(2))

            # Save the modified document as a Word document
            modified_folder = 'E:/fyp/static/modified/'
            timestamp = time.strftime("%H%M%S")
            modified_path_docx =('modified_document.docx')
            doc.save(modified_path_docx)

            # Convert the Word document to PDF
            modified_path_pdf = os.path.join(modified_folder, f'modified_document_{timestamp}.pdf')
            convert(modified_path_docx, modified_path_pdf)
            modified_pdf=(f'modified_document_{timestamp}.pdf')

            update_query = "UPDATE fyp_title SET digital_signature_link = %s WHERE id = %s"
            mydb.cursor().execute(update_query, (modified_pdf, staff_id))
            mydb.commit()

            # Return the converted PDF for download
            return send_from_directory(modified_folder, modified_pdf, as_attachment=True)
            # Return the converted PDF for download
            return send_from_directory(modified_folder, modified_pdf, as_attachment=True)

        finally:
            
            pythoncom.CoUninitialize()

    # Handle the case where data1 and data2 are not available
    return render_template('download_img.html', staff_id=staff_id, data1=data1, staffid=staffid, data2=data2, msg=msg)



@app.route('/view_report/<filename>')
def view_report(filename):
    # Assuming filename is the name of the DOCX, PDF, or TXT file
    file_path = os.path.join('static/uploads', filename)

    try:
        if filename.endswith('.docx'):
            # Read the DOCX file and extract text
            doc = Document(file_path)
            text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])

        elif filename.endswith('.pdf'):
            # Read the PDF file and extract text
            with fitz.open(file_path) as pdf_document:
                text_content = ""
                for page_num in range(pdf_document.page_count):
                    page = pdf_document[page_num]
                    text_content += page.get_text("text")

        elif filename.endswith('.txt'):
            # Read the TXT file
            with open(file_path, 'r') as txt_file:
                text_content = txt_file.read()

        else:
            text_content = "Unsupported file format. Only DOCX, PDF, and TXT files are supported."

        # Convert the text content to PDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, text_content)
        pdf_output_path = os.path.join('static/uploads', f'{filename}.pdf')
        pdf.output(pdf_output_path)

        # Return the PDF file path to be displayed in the browser
        return render_template('view_report.html', pdf_path=f'/static/uploads/{filename}.pdf')

    except Exception as e:
        # Handle any errors that might occur during text extraction or PDF conversion
        print(f"Error processing {filename}: {str(e)}")
        error_message = "Error processing the document."
        return render_template('view_report.html', error_message=error_message)



@app.route('/view_report12/<filename>')
def view_report12(filename):
    # Assuming filename is the name of the DOCX, PDF, or TXT file
    file_path = os.path.join('static/report', filename)

    try:
        if filename.endswith('.docx'):
            # Read the DOCX file and extract text
            doc = Document(file_path)
            text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])

        elif filename.endswith('.pdf'):
            # Read the PDF file and extract text
            with fitz.open(file_path) as pdf_document:
                text_content = ""
                for page_num in range(pdf_document.page_count):
                    page = pdf_document[page_num]
                    text_content += page.get_text("text")

        elif filename.endswith('.txt'):
            # Read the TXT file
            with open(file_path, 'r') as txt_file:
                text_content = txt_file.read()

        else:
            text_content = "Unsupported file format. Only DOCX, PDF, and TXT files are supported."

        # Convert the text content to PDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, text_content)
        pdf_output_path = os.path.join('static/report', f'{filename}.pdf')
        pdf.output(pdf_output_path)

        # Return the PDF file path to be displayed in the browser
        return render_template('view_report12.html', pdf_path=f'/static/report/{filename}.pdf')

    except Exception as e:
        # Handle any errors that might occur during text extraction or PDF conversion
        print(f"Error processing {filename}: {str(e)}")
        error_message = "Error processing the document."
        return render_template('view_report12.html', error_message=error_message)






####################################################################################################3


@app.route("/call", methods=["GET", "POST"])
def call():

    aid=request.args.get("aid")
    if request.method == "POST":
        room_id = request.form['room_id']
        cursor = mydb.cursor()
        cursor.execute("update fyp_review set link=%s where id=%s",(room_id, aid))
        mydb.commit()
        
        return redirect(url_for("entry_checkpoint", room_id=room_id, aid=aid))

    return render_template("call.html")

@app.route("/room/<string:room_id>/")
def enter_room(room_id):
    act=request.args.get("act")
    
    
    if room_id not in session:
        return redirect(url_for("entry_checkpoint", room_id=room_id))
    
    return render_template("chatroom.html", room_id=room_id, display_name=session[room_id]["name"], mute_audio=session[room_id]["mute_audio"], mute_video=session[room_id]["mute_video"])

@app.route("/room/<string:room_id>/checkpoint/", methods=["GET", "POST"])
def entry_checkpoint(room_id):
    

    username=""
    
    if request.method == "POST":
        mute_audio = request.form['mute_audio']
        mute_video = request.form['mute_video']
        session[room_id] = {"name": username, "mute_audio":mute_audio, "mute_video":mute_video}
        return redirect(url_for("enter_room", room_id=room_id))

    return render_template("chatroom_checkpoint.html", room_id=room_id)

@socketio.on("connect")
def on_connect():
    sid = request.sid
    print("New socket connected ", sid)
    

@socketio.on("join-room")
def on_join_room(data):
    sid = request.sid
    room_id = data["room_id"]
    display_name = session[room_id]["name"]
    
    # register sid to the room
    join_room(room_id)
    _room_of_sid[sid] = room_id
    _name_of_sid[sid] = display_name
    
    # broadcast to others in the room
    print("[{}] New member joined: {}<{}>".format(room_id, display_name, sid))
    emit("user-connect", {"sid": sid, "name": display_name}, broadcast=True, include_self=False, room=room_id)
    
    # add to user list maintained on server
    if room_id not in _users_in_room:
        _users_in_room[room_id] = [sid]
        emit("user-list", {"my_id": sid}) # send own id only
    else:
        usrlist = {u_id:_name_of_sid[u_id] for u_id in _users_in_room[room_id]}
        emit("user-list", {"list": usrlist, "my_id": sid}) # send list of existing users to the new member
        _users_in_room[room_id].append(sid) # add new member to user list maintained on server

    print("\nusers: ", _users_in_room, "\n")


@socketio.on("disconnect")
def on_disconnect():
    sid = request.sid
    room_id = _room_of_sid[sid]
    display_name = _name_of_sid[sid]

    print("[{}] Member left: {}<{}>".format(room_id, display_name, sid))
    emit("user-disconnect", {"sid": sid}, broadcast=True, include_self=False, room=room_id)

    _users_in_room[room_id].remove(sid)
    if len(_users_in_room[room_id]) == 0:
        _users_in_room.pop(room_id)

    _room_of_sid.pop(sid)
    _name_of_sid.pop(sid)

    print("\nusers: ", _users_in_room, "\n")


@socketio.on("data")
def on_data(data):
    sender_sid = data['sender_id']
    target_sid = data['target_id']
    if sender_sid != request.sid:
        print("[Not supposed to happen!] request.sid and sender_id don't match!!!")

    if data["type"] != "new-ice-candidate":
        print('{} message from {} to {}'.format(data["type"], sender_sid, target_sid))
    socketio.emit('data', data, room=target_sid)





#FORUM########################################################################################3




   
@app.route('/logout')
def logout():
    session.clear()
    print("Logged out successfully", 'success')
    return redirect(url_for('index'))



if __name__ == "__main__":
    app.secret_key = os.urandom(12)
    socketio.run(app, debug=True, host='0.0.0.0', port=5000)

